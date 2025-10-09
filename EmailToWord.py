import os
import re
import shutil
import tempfile
import unicodedata
from typing import Optional, List, Dict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

import msg_parser  # fallback
from bs4 import BeautifulSoup

# =========================
# ===== USER SETTINGS =====
# =========================
EMAIL_FOLDER   = r"C:\Users\Avinash\Downloads\emails"
OUTPUT_DOCX    = "compiled_emails.docx"
CAUTION_PHRASE = "Classification: Wisayah: Company General Use"

STOP_PHRASES = [
    "You have received this email because you are subscribed", "For important information including analyst certification and disclosures",
    "Research Terms and Disclosures", "Morgan Stanley is not acting as a municipal advisor", "This e-mail and any attachments",
    "Disclaimer", "Disclosure", "Guide to Analysis", "Remove me from this distribution",
]
MAX_IMG_WIDTH_IN = 6.0

# ==================================
# ===== CORE EMAIL PROCESSING ======
# ==================================

def _to_text(x):
    if x is None: return None
    if isinstance(x, bytes):
        for enc in ("utf-8", "utf-16", "cp1252", "latin-1"):
            try: return x.decode(enc)
            except Exception: pass
        return x.decode("utf-8", errors="ignore")
    return str(x)

def _compile_flexible_phrase_regex(phrase: str) -> re.Pattern:
    norm = unicodedata.normalize("NFKC", phrase)
    norm = norm.replace("⚠️", "(?:⚠️|⚠|&#9888;|&#x26A0;)").replace("⚠", "(?:⚠️|⚠|&#9888;|&#x26A0;)")
    tokens = [re.escape(t) for t in re.split(r"\s+", norm.strip())]
    pattern = r"(?:\s|<[^>]*>)*".join(tokens)
    return re.compile(pattern, re.IGNORECASE | re.DOTALL)

def _first_stop_index(content: str, stop_phrases: List[str]) -> Optional[int]:
    earliest = None
    for phrase in stop_phrases:
        rx = _compile_flexible_phrase_regex(phrase)
        m = rx.search(content)
        if m and (earliest is None or m.start() < earliest):
            earliest = m.start()
    return earliest

def _cut_by_markers(body: str, start_phrase: str, stop_phrases: List[str]) -> Optional[str]:
    start_re = _compile_flexible_phrase_regex(start_phrase)
    matches = list(start_re.finditer(body))
    if not matches:
        return None # Start phrase was not found
    content = body[matches[-1].end():]
    cut_at = _first_stop_index(content, stop_phrases)
    if cut_at is not None:
        content = content[:cut_at]
    return content.strip() or None

def extract_email_bodies(email_path: str):
    html_body, text_body = None, None
    try:
        import extract_msg
        m = extract_msg.Message(email_path)
        html_body = _to_text(getattr(m, "htmlBody", None))
        text_body = _to_text(getattr(m, "body", None))
    except Exception as e:
        print(f"      -> [extract_msg error]: {e}")
        try:
            msg = msg_parser.MsOxMessage(email_path)
            html_body = _to_text(getattr(msg, "html_body", None))
            text_body = _to_text(getattr(msg, "body", None))
        except Exception as e2:
            print(f"      -> [msg_parser error]: {e2}")
    return html_body, text_body

# =================================
# ===== FINAL FORMATTING LOGIC ====
# =================================

def final_formatter(filepath: str):
    """
    Rebuilds the document by grouping consecutive text blocks and ensuring
    single-blank-line spacing between these groups and tables.
    """
    print(f"--- Applying final formatting to '{filepath}' ---")
    try:
        source_doc = Document(filepath)
        if not source_doc.paragraphs:
            print("     -> Document is empty, skipping formatting.")
            return

        new_doc = Document()

        for style in source_doc.styles:
            if style.name not in [s.name for s in new_doc.styles]:
                try: new_doc.styles.add_style(style.name, style.type)
                except Exception: pass

        groups = []
        current_para_group = []
        for element in source_doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, source_doc)
                if para.text.strip():
                    current_para_group.append(para)
                else:
                    if current_para_group:
                        groups.append(current_para_group)
                        current_para_group = []
            elif isinstance(element, CT_Tbl):
                if current_para_group:
                    groups.append(current_para_group)
                    current_para_group = []
                groups.append(Table(element, source_doc))
        if current_para_group:
            groups.append(current_para_group)

        for i, group in enumerate(groups):
            if isinstance(group, Table):
                new_doc.element.body.append(group._element)
            elif isinstance(group, list):
                for para in group:
                    new_para = new_doc.add_paragraph(style=para.style.name)
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold, new_run.italic, new_run.underline = run.bold, run.italic, run.underline
                        if run.font.name: new_run.font.name = run.font.name
                        if run.font.size: new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
            
            if i < len(groups) - 1:
                new_doc.add_paragraph()

        new_doc.save(filepath)
        print("✅ Final formatting applied successfully.")

    except Exception as e:
        print(f"❌ An error occurred during final formatting: {e}")

# =========================
# ========= MAIN ==========
# =========================
def main():
    if not os.path.isdir(EMAIL_FOLDER):
        print(f"Error: The folder '{EMAIL_FOLDER}' does not exist.")
        return

    doc = Document()
    doc.add_heading('Compiled Email Content', level=1)

    email_files = sorted([f for f in os.listdir(EMAIL_FOLDER) if f.lower().endswith(".msg")])
    if not email_files:
        print("No .msg files found.")
        return

    print(f"Found {len(email_files)} .msg files to process.")
    
    for filename in email_files:
        email_path = os.path.join(EMAIL_FOLDER, filename)
        print(f"\n--- Processing: {filename} ---")

        html_body, text_body = extract_email_bodies(email_path)
        
        # We must search in the HTML body if it exists, otherwise use the text body.
        body_to_search = html_body or text_body
        
        if not body_to_search:
            print("     -> Could not extract any body content. Skipping.")
            continue
            
        cut_content = _cut_by_markers(body_to_search, CAUTION_PHRASE, STOP_PHRASES)

        if not cut_content:
            print("     -> Content found but was empty after cutting headers/footers. Skipping.")
            continue

        # Convert the isolated content to clean plain text.
        soup = BeautifulSoup(cut_content, "lxml")
        final_text = soup.get_text(separator='\n', strip=True)
        
        print("     -> Content extracted and cleaned.")
        doc.add_heading(f"Content from: {filename}", level=2)
        
        # Add the text to the document. The formatter will fix the spacing.
        doc.add_paragraph(final_text)
        doc.add_page_break()

    print(f"\n✅ Initial document compiled: {os.path.abspath(OUTPUT_DOCX)}")
    doc.save(OUTPUT_DOCX)
    
    final_formatter(OUTPUT_DOCX)


if __name__ == "__main__":
    # Ensure necessary libraries are installed:
    # pip install python-docx extract-msg msg-parser beautifulsoup4 lxml
    main()