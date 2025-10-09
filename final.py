import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph

def fix_paragraph_spacing(filepath: str):
    """
    Opens a .docx file, ensures a single empty line between text paragraphs,
    and preserves tables. Overwrites the original file.

    Args:
        filepath: The path to the .docx file to format.
    """
    print(f"--- Starting final formatting for '{filepath}' ---")
    
    try:
        source_doc = Document(filepath)
        new_doc = Document()

        # Copy styles from the old document to the new one to preserve formatting.
        for style in source_doc.styles:
            if style.name not in [s.name for s in new_doc.styles]:
                try:
                    new_doc.styles.add_style(style.name, style.type)
                except Exception:
                    pass

        # This flag tracks if the very last thing we added was a paragraph with text.
        last_element_was_text_para = False

        # Iterate through the document's body elements (paragraphs and tables).
        for element in source_doc.element.body:
            # --- Case 1: The element is a Paragraph ---
            if isinstance(element, CT_P):
                para = Paragraph(element, source_doc)

                # Check if the paragraph contains any visible text.
                if para.text.strip():
                    # If the last element was also a text paragraph, we must add a blank line.
                    if last_element_was_text_para:
                        new_doc.add_paragraph()

                    # **THE FIX IS HERE**: Use the style NAME, not the style object.
                    new_para = new_doc.add_paragraph(style=para.style.name)
                    
                    # Copy the text and formatting run by run.
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        # Copy basic run formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                    
                    last_element_was_text_para = True
                
                else:
                    # If the paragraph is blank, it resets our flag.
                    last_element_was_text_para = False

            # --- Case 2: The element is a Table ---
            elif isinstance(element, CT_Tbl):
                # A table resets the flag.
                last_element_was_text_para = False
                
                # Append the entire table element to the new document's body.
                new_doc.element.body.append(element)
        
        # Save the new, formatted document, overwriting the old one.
        new_doc.save(filepath)
        print(f"✅ Successfully formatted paragraphs in '{filepath}'")

    except Exception as e:
        print(f"❌ An error occurred during formatting: {e}")
        print("   Please ensure the file is not open elsewhere and that it is a valid .docx file.")


if __name__ == "__main__":
    file_to_fix = "compiled_emails.docx"
    
    if os.path.exists(file_to_fix):
        fix_paragraph_spacing(file_to_fix)
    else:
        print(f"Error: The file '{file_to_fix}' was not found in this directory.")