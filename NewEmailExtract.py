from PIL import Image, ImageChops
import io
import os
import sqlite3
import json
import tempfile
import re
from datetime import datetime, timedelta
import uuid
import io
import base64
import pypandoc

import streamlit as st
import pandas as pd
from extract_msg import Message
from openai import OpenAI
import requests
import msal
from thefuzz import fuzz, process
from sqlalchemy import create_engine, text
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
import docx
from bs4 import BeautifulSoup
import unicodedata
from docx.shared import Inches
from bs4.element import Tag, NavigableString

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="Email Intelligence Extractor",
    page_icon="📧",
    layout="wide"
)

# --- STYLING ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="st-"] {
        font-family: 'Inter', sans-serif;
        color: #262626; /* Darker text for better contrast */
    }
    
    .st-emotion-cache-1y4p8pa { /* Main content area */
        max-width: 100%;
        padding: 2.5rem 4rem; /* Increased padding */
    }

    .header-container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding-bottom: 1.5rem;
        border-bottom: 1px solid #e0e0e0;
        margin-bottom: 2.5rem;
    }
    
    .app-title {
        font-size: 2.5em;
        font-weight: 700;
        color: #1a1a1a;
        letter-spacing: -0.02em;
    }
    
    .stButton>button {
        border-radius: 8px;
        border: 1px solid #d1d5db;
        background-color: #ffffff;
        color: #374151;
        font-weight: 500;
        transition: all 0.2s ease;
        padding: 0.6rem 1.2rem;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    
    .stButton>button:hover {
        background-color: #f0f2f6;
        border-color: #a0a4ac;
        color: #1a1a1a;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    button[data-testid="stButton-primary"] {
        background-color: #1e70bf !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.2rem !important;
        font-weight: 500 !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
        transition: all 0.2s ease !important;
    }

    button[data-testid="stButton-primary"]:hover {
        background-color: #155a9b !important;
        color: white !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.15) !important;
    }

    .st-emotion-cache-16txtl3 {
        padding: 2rem;
        background-color: #ffffff;
        border-radius: 12px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 1.5rem;
    }

    h1, h2, h3, h4, h5, h6 {
        font-weight: 600;
        color: #1a1a1a;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
    }
    h2 { font-size: 1.8em; }
    h3 { font-size: 1.3em; margin-top: 1.2rem; }

    .stTextInput>div>div>input {
        border-radius: 8px;
        border: 1px solid #d1d5db;
        padding: 0.7rem 1rem;
        font-size: 1rem;
        box-shadow: inset 0 1px 2px rgba(0,0,0,0.03);
    }
    .stTextInput>div>div>input:focus {
        border-color: #1e70bf;
        box-shadow: 0 0 0 2px rgba(30, 112, 191, 0.2);
        outline: none;
    }

    .stTabs [data-baseweb="tab-list"] button {
        background-color: #f8f8f8;
        color: #666666;
        border-radius: 8px 8px 0 0;
        padding: 0.7rem 1.2rem;
        font-weight: 500;
        border: none;
        margin-right: 5px;
        transition: all 0.2s ease;
    }
    .stTabs [data-baseweb="tab-list"] button:hover {
        background-color: #e8e8e8;
        color: #333333;
    }
    .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {
        background-color: #ffffff;
        color: #1e70bf;
        border-bottom: 3px solid #1e70bf;
        font-weight: 600;
        box-shadow: none;
    }
    .stTabs [data-baseweb="tab-panel"] {
        padding-top: 1.5rem;
        border-top: 1px solid #e0e0e0;
        margin-top: -1px;
    }
</style>
""", unsafe_allow_html=True)


# --- CONFIGURATION & CLIENTS ---
try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    AZURE_CLIENT_ID = st.secrets["AZURE_CLIENT_ID"]
    AZURE_TENANT_ID = st.secrets["AZURE_TENANT_ID"]
    AZURE_CLIENT_SECRET = st.secrets["AZURE_CLIENT_SECRET"]
    DB_CONNECTION_STRING = st.secrets["DB_CONNECTION_STRING"]
    
    connect_str = st.secrets["AZURE_STORAGE_CONNECTION_STRING"]
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    AZURE_CONTAINER_NAME = st.secrets["AZURE_CONTAINER_NAME"]

except KeyError as e:
    st.error(f"Configuration key not found in Streamlit Secrets: {e}")
    st.stop()

# --- CONSTANTS ---
MASTER_DB_NAME = "Master_Company_List.db"
AUTHORITY = f"https://login.microsoftonline.com/{AZURE_TENANT_ID}"
SCOPE = ["https://graph.microsoft.com/.default"]

# --- DATABASE SETUP ---
engine = create_engine(DB_CONNECTION_STRING)

def query_db(query, params=None):
    with engine.connect() as conn:
        return pd.read_sql_query(sql=text(query), con=conn, params=params)

def query_local_db(db_name, query, params=None):
    with sqlite3.connect(db_name) as conn:
        return pd.read_sql_query(query, conn, params=params if params else ())

@st.cache_data(ttl=3600)
def get_master_company_data():
    if not os.path.exists(MASTER_DB_NAME):
        return pd.DataFrame()
    return query_local_db(MASTER_DB_NAME, "SELECT * FROM master_companies")

@st.cache_data(ttl=3600)
def get_master_broker_names():
    if not os.path.exists(MASTER_DB_NAME): 
        return []
    df = query_local_db(MASTER_DB_NAME, "SELECT Name FROM master_brokers")
    return df['Name'].tolist() if not df.empty else []

def insert_into_db(data):
    data_lowercase_keys = {key.lower(): value for key, value in data.items()}
    df = pd.DataFrame([data_lowercase_keys])

    # 1. Cast types to handle 'None' as a number
    if 'fiscalyear' in df.columns:
        df['fiscalyear'] = df['fiscalyear'].astype('Int64')
    if 'fiscalquarter' in df.columns:
        df['fiscalquarter'] = df['fiscalquarter'].astype('Int64')

    # --- THIS IS THE FIX ---
    # 2. Clean null characters ('\x00' or '\u0000') from text fields.
    #    This is what your new error log is complaining about.
    if 'emailcontent' in df.columns and df['emailcontent'].iloc[0] is not None:
        df['emailcontent'] = df['emailcontent'].astype(str).str.replace(r'\x00|\u0000', '', regex=True)
    if 'emailsubject' in df.columns and df['emailsubject'].iloc[0] is not None:
        df['emailsubject'] = df['emailsubject'].astype(str).str.replace(r'\x00|\u0000', '', regex=True)
    # --- END FIX ---

    # 3. Convert all pd.NA/np.nan to Python 'None' for the driver
    df = df.astype(object).where(pd.notnull(df), None)

    with engine.connect() as conn:
        df.to_sql('email_data', con=conn, if_exists='append', index=False)
        conn.commit()

# --- MATCHING LOGIC ---
def normalize_company_name(name):
    if not isinstance(name, str):
        return ""
    name = name.lower()
    suffixes = [
        r'\bholding\b', r'\bholdings\b', r'\bhold\b', r'\bltd\b', r'\binc\b',
        r'\bcorp\b', r'\bcorporation\b', r'\bgroup\b', r'\bplc\b', r'\bco\b'
    ]
    for suffix in suffixes:
        name = re.sub(suffix, '', name, flags=re.IGNORECASE)
    name = re.sub(r'[^\w\s]', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name

def find_company_in_master(extracted_report, master_df):
    company_name = extracted_report.get("Company")
    ticker = extracted_report.get("Ticker")
    name_columns = ['short_name', 'company_short_name', 'full_company_name', 'acronym']
    
    for col in name_columns:
        if col not in master_df.columns:
            master_df[col] = None

    if ticker and isinstance(ticker, str):
        mask = master_df['ticker'].str.lower() == ticker.lower()
        match = master_df[mask.fillna(False)]
        if not match.empty:
            return match.iloc[0], "Ticker Match"

    if not company_name or not isinstance(company_name, str):
        return None, "No Match (Invalid/Missing Company Name)"

    for col in name_columns:
        if pd.api.types.is_string_dtype(master_df[col]):
            match = master_df[master_df[col].str.lower() == company_name.lower()]
            if not match.empty:
                return match.iloc[0], f"Exact Match ({col})"

    normalized_input = normalize_company_name(company_name)
    if not normalized_input:
        return None, "No Match (Empty Normalized Name)"

    normalized_cols = []
    for col in name_columns:
        normalized_col_name = f'normalized_{col}'
        if normalized_col_name not in master_df.columns:
            master_df[normalized_col_name] = master_df[col].fillna('').apply(normalize_company_name)
        normalized_cols.append(normalized_col_name)

    for norm_col in normalized_cols:
        match = master_df[master_df[norm_col] == normalized_input]
        if not match.empty:
            original_col = norm_col.replace('normalized_', '')
            return match.iloc[0], f"Normalized Match ({original_col})"

    all_substring_matches = pd.DataFrame()
    for norm_col in normalized_cols:
        valid_rows = master_df[master_df[norm_col] != '']
        substring_matches = valid_rows[valid_rows[norm_col].str.contains(normalized_input, na=False)]
        if not substring_matches.empty:
            all_substring_matches = pd.concat([all_substring_matches, substring_matches])
    
    all_substring_matches = all_substring_matches.drop_duplicates()
    if len(all_substring_matches) == 1:
        return all_substring_matches.iloc[0], "Substring Match"

    def get_max_fuzzy_score(row):
        scores = [fuzz.token_set_ratio(company_name, str(row[col])) for col in name_columns if pd.notna(row[col])]
        return max(scores) if scores else 0

    master_df['fuzzy_score'] = master_df.apply(get_max_fuzzy_score, axis=1)
    
    if not master_df.empty:
        best_match = master_df.loc[master_df['fuzzy_score'].idxmax()]
        if best_match['fuzzy_score'] >= 95:
            return best_match, f"Fuzzy Match ({best_match['fuzzy_score']}%)"

    return None, "No Match"

def find_broker_in_master(extracted_broker_name, master_broker_list):
    if not extracted_broker_name or not master_broker_list:
        return "Unknown", 0
    best_match, score = process.extractOne(extracted_broker_name, master_broker_list)
    return (best_match, score) if score >= 85 else (extracted_broker_name, score)

# --- OUTLOOK & PARSING LOGIC ---
@st.cache_resource(ttl=3500)
def get_graph_api_token():
    app = msal.ConfidentialClientApplication(client_id=AZURE_CLIENT_ID, authority=AUTHORITY, client_credential=AZURE_CLIENT_SECRET)
    result = app.acquire_token_silent(scopes=SCOPE, account=None) or app.acquire_token_for_client(scopes=SCOPE)
    if "access_token" in result: return result['access_token']
    st.error("Failed to acquire access token."); st.json(result.get("error_description"))
    return None

def scan_outlook_emails(user_id, token, sender_domain=None):
    headers = {'Authorization': f'Bearer {token}'}
    query_filter = "isRead eq false"
    if sender_domain: query_filter += f" and contains(from/emailAddress/address, '{sender_domain}')"
    endpoint = f"https://graph.microsoft.com/v1.0/users/{user_id}/mailFolders/inbox/messages"
    params = {'$filter': query_filter, '$select': 'subject,body,from', '$top': '25'}
    try:
        response = requests.get(endpoint, headers=headers, params=params)
        response.raise_for_status()
        return response.json().get('value', [])
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching emails: {e}"); st.json(e.response.json())
        return None


# ⭐️ --- NEW CLEANING FUNCTIONS TO ADD --- ⭐️

def clean_plain_text_for_llm(text):
    """
    Aggressively cleans plain text *before* sending to the LLM.
    Removes forwarded headers, classification junk, etc.
    """
    if not text:
        return ""
    
    # Remove forwarded headers (multi-line)
    text = re.sub(r'From:.*Sent:.*To:.*Subject:.*', '', text, flags=re.DOTALL | re.IGNORECASE)
    # Remove Wisayah classification junk
    text = re.sub(r'This is classified as Wisayah.*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Classification: Wisayah.*', '', text, flags=re.IGNORECASE)
    # Remove the "Just text" artifact
    text = re.sub(r'^\s*Just text\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
    # Remove the Outlook "don't often get email" warning
    text = re.sub(r"Some people who received this message don't often get email from.*Learn why.*", '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Strip leading/trailing whitespace
    return text.strip()

def clean_html_for_database(html_string):
    """
    Pass-through function.
    We will store the raw HTML and do all cleaning during the export step.
    This ensures data is saved reliably.
    """
    if not html_string:
        return ""
    # Store the raw HTML as-is to ensure the database insert never fails
    return str(html_string)



def extract_info_with_chatgpt(subject, body, master_brokers, email_theme=None):
    broker_list_str = ", ".join(master_brokers)
    
    # NEW: Add context if a theme is provided
    theme_context = ""
    if email_theme and email_theme != "Select a category...":
        theme_context = f"The user has pre-categorized this email with the theme: '{email_theme}'. Use this as a strong hint."

    prompt = f"""You are an expert financial analyst. From the email below, extract key details for the financial report.
    
    {theme_context}

    **Instructions:**
    1.  **Report Type:** First, determine if this is a report about a *specific company* OR if it is a *thematic/macro* report (e.g., 'Global Credit Strategy', 'Macro Round-up', 'EM Strategy').
    2.  **Company:** If it's about a *specific company*, extract its name. If it is a *thematic/macro* report, set the 'Company' to a descriptive title like 'Macro Report', 'Thematic Report', or 'Strategy Note'.
    3.  **Ticker:** Extract the stock Ticker if it is a specific company report. If it is a *thematic/macro* report, set 'Ticker' to null.
    4.  **BrokerName:** You MUST choose the most appropriate name from this list of known brokers: {broker_list_str}. If no suitable broker is mentioned or found, classify it as 'Unknown'.
    5.  **Category:** High-level classification like 'Equity Research', 'Macro Research', 'Credit Strategy'.
    6.  **ContentType:** Must be from this specific list: 'Earnings Commentary', 'Earnings Call Commentary', 'Market Update', 'Stock Initiation', 'Thematic Note', 'Strategy Note', 'Macro Note', 'Other'.
    7.  **FiscalYear:** The four-digit year the report is about (e.g., 2024, 2025). Look for terms like '4Q24', 'FY25', or '2025 Results'. If not explicitly stated, infer from the email's content. If it cannot be determined, return null.
    8.  **FiscalQuarter:** The quarter the report is about as a single number (1, 2, 3, or 4). Look for terms like 'Q3', '3Q', 'Third Quarter'. If not explicitly stated, infer from the content. If it cannot be determined, return null.

    **Email Details:**
    - EMAIL SUBJECT: {subject}
    - EMAIL BODY (first 8000 characters):
    ---
    {body[:8000]}
    ---
    
    Provide the output in a JSON object with a single key "reports", which is a list. Even if it's a macro report, it should be one entry in the list.
    Example (Company Report): {{"reports": [{{"Country": "USA", "Sector": "Technology", "Company": "Example Corp", "Ticker": "EXMPL", "Category": "Equity Research", "ContentType": "Earnings Commentary", "BrokerName": "Global Brokerage", "FiscalYear": 2025, "FiscalQuarter": 3}}]}}
    Example (Macro Report): {{"reports": [{{"Country": "Global", "Sector": "Macro", "Company": "Macro Report", "Ticker": null, "Category": "Macro Research", "ContentType": "Macro Note", "BrokerName": "Global Brokerage", "FiscalYear": null, "FiscalQuarter": null}}]}}"""
    try:
        response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": "You are a helpful assistant designed to output JSON."}, {"role": "user", "content": prompt}], response_format={"type": "json_object"})
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"Error with OpenAI API call: {e}"); return None

# --- CORE PROCESSING FUNCTION ---
# ⭐️ CHANGE: Added 'email_theme=None' parameter
def process_emails(email_source, source_type, email_theme=None):
    master_companies_df = get_master_company_data()
    master_brokers = get_master_broker_names()

    if master_companies_df.empty:
        st.error(f"Master company database '{MASTER_DB_NAME}' is empty or not found.")
        return

    status_container = st.container() 
    progress_bar = st.progress(0, text="Initializing...")
    total_emails = len(email_source)
    
    for i, item in enumerate(email_source):
        subject, plain_body, raw_html_body = (None, None, None)
        blob_name = None

        try:
            if source_type == 'upload':
                file_bytes = item.getvalue()
                blob_name = f"emails/{uuid.uuid4()}-{item.name}"
                blob_client = blob_service_client.get_blob_client(container=AZURE_CONTAINER_NAME, blob=blob_name)
                blob_client.upload_blob(file_bytes)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".msg") as tmp:
                    tmp.write(file_bytes)
                
                with Message(tmp.name) as msg:
                    subject = msg.subject
                    # ⭐️ CLEAN THE PLAIN TEXT BEFORE SENDING TO LLM
                    plain_body = clean_plain_text_for_llm(msg.body)
                    
                    # === IMPROVED HTML BODY PROCESSING ===
                    html_body_content = msg.htmlBody
                    
                    if isinstance(html_body_content, bytes):
                        try:
                            # Attempt to decode the byte string, UTF-8 is most common
                            raw_html_body = html_body_content.decode('utf-8')
                        except UnicodeDecodeError:
                            # If UTF-8 fails, fallback to a more lenient encoding like latin-1
                            raw_html_body = html_body_content.decode('latin-1', errors='ignore')
                    else:
                        # It's already a string or None, so use it as is
                        raw_html_body = html_body_content
                    
                    # === Handle hexadecimal encoded HTML content ===
                    if raw_html_body and raw_html_body.startswith('\\x'):
                        try:
                            # Remove the '\\x' prefix and decode the hex string
                            hex_string = raw_html_body.replace('\\x', '')
                            decoded_bytes = bytes.fromhex(hex_string)
                            raw_html_body = decoded_bytes.decode('utf-8', errors='ignore')
                            status_container.success(f"✅ Successfully decoded hexadecimal content for: {subject}")
                        except Exception as hex_decode_error:
                            status_container.warning(f"Failed to decode hexadecimal content for {subject}: {hex_decode_error}")
                            # Fallback to plain text if hex decoding fails
                            raw_html_body = f"<pre>{plain_body}</pre>"
                    
                    # Final fallback to plain text if HTML body is still not available
                    if not raw_html_body:
                         raw_html_body = f"<pre>{plain_body}</pre>"
                    # === END OF FIX ===

                os.unlink(tmp.name)
            
            elif source_type == 'outlook':
                subject = item.get('subject', 'No Subject')
                # ⭐️ CLEAN THE PLAIN TEXT BEFORE SENDING TO LLM
                plain_body = clean_plain_text_for_llm(item.get('body', {}).get('content', ''))
                raw_html_body = item.get('body', {}).get('content', '')

            # ⭐️ CLEAN THE HTML *BEFORE* SAVING TO DATABASE
            raw_html_body = clean_html_for_database(raw_html_body)
                
            if blob_name:
                status_container.info(f"📤 Saved original email")

        except Exception as e:
            st.error(f"Failed during file handling or Azure upload: {e}")
            continue

        progress_bar.progress((i + 1) / total_emails, text=f"Processing: {subject}")
        if not (subject and plain_body and raw_html_body):
            status_container.warning(f"Skipping an email due to parsing error.")
            continue
        
        # ⭐️ Use the *already-cleaned* plain_body for the LLM
        
        extracted = extract_info_with_chatgpt(subject, plain_body, master_brokers, email_theme)

        # --- START OF NEW DEBUG CODE ---
        if not extracted or "reports" not in extracted or not extracted["reports"]:
            status_container.warning(f"⚠️ OpenAI failed to extract any reports from: '{subject}'. Nothing will be saved for this email.")
            continue # Skip to the next email
        # --- END OF NEW DEBUG CODE ---
        if not (extracted and "reports" in extracted): continue

        for report in extracted["reports"]:
            report.setdefault('Company', 'N/A')
            report.setdefault('Ticker', None)
            report.setdefault('BrokerName', 'Unknown')
            report.setdefault('FiscalYear', None)
            report.setdefault('FiscalQuarter', None)
            report.setdefault('Category', 'N/A')
            report.setdefault('ContentType', 'Other')
            report.setdefault('Country', 'N/A')
            report.setdefault('Sector', 'N/A')
            
            report['EmailSubject'] = subject
            report['EmailContent'] = raw_html_body  # This should now be properly decoded
            report['blob_name'] = blob_name

            # ⭐️ CHANGE: Add the selected email theme to the report data
            # If the theme is "Select a category...", store None (NULL)
            if email_theme and email_theme != "Select a category...":
                report['EmailTheme'] = email_theme
            else:
                report['EmailTheme'] = None # Explicitly set to None for clarity

            company_to_find = report.get("Company", "N/A")
            matched_row, match_status = find_company_in_master(report, master_companies_df.copy())

            if matched_row is not None:
                report["Company"] = matched_row['short_name']
                report["Ticker"] = matched_row['ticker']
                report["Country"] = matched_row['country']
                report["Sector"] = matched_row['sector']
            else:
                status_container.warning(f"❌ Could not find a match for '{company_to_find}'")

            report["MatchStatus"] = match_status
            # --- START OF NEW DEBUG CODE ---
            try:
                status_container.success(f"✅ Extracted report for '{report.get('Company', 'N/A')}' from '{subject}'. Attempting to save...")
                insert_into_db(report)
            except Exception as db_error:
                status_container.error(f"❌ DATABASE ERROR for '{subject}': {db_error}")
                st.error(f"Failed to insert data: {report}")
            # --- END OF NEW DEBUG CODE ---

    progress_bar.progress(1.0, text="Processing complete!")
    st.success("✅ Processing complete! The database has been updated.")

# --- AZURE SAS URL HELPER ---
def generate_sas_url(container_name, blob_name):
    if not blob_name or pd.isna(blob_name):
        return None
    try:
        account_name = blob_service_client.account_name
        account_key = blob_service_client.credential.account_key
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container_name,
            blob_name=blob_name,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=1)
        )
        return f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_name}?{sas_token}"
    except Exception:
        return None

# ⭐️ --- NEW AGGRESSIVE HTML CLEANING FUNCTION --- ⭐️
# ⭐️ --- NEW AGGRESSIVE HTML CLEANING FUNCTION --- ⭐️
# ⭐️ --- NEW ALL-IN-ONE EXPORT CLEANING FUNCTION --- ⭐️
def clean_html_for_export(html_string, temp_dir, msg_file_path=None):
    """
    Aggressively cleans HTML from the database for Word export.
    This function removes ALL junk text, disclaimers, and signatures,
    and also processes images and fixes layout for Pandoc.
    """
    if not html_string:
        return ""
        
    soup = BeautifulSoup(html_string, 'html.parser')
    
    # --- Define patterns for removal ---
    removal_patterns = [
        # --- Classification/Warnings ---
        re.compile(r'Classification: Wisayah', re.I),
        re.compile(r'This is classified as Wisayah', re.I),
        re.compile(r'CAUTION: This email has been sent from outside', re.I),
        re.compile(r"don't often get email from", re.I),
        
        # --- Junk text artifacts ---
        re.compile(r'^\s*Just text\s*$', re.I),
        re.compile(r'^\s*The following table:\s*$', re.I),
        
        # --- Email Headers (as individual lines) ---
        re.compile(r'^\s*From:', re.I),
        re.compile(r'^\s*Sent:', re.I),
        re.compile(r'^\s*To:', re.I),
        re.compile(r'^\s*Subject:', re.I),

        # --- Signatures / Analyst Names ---
        re.compile(r'Eric Robertsen', re.I),
        re.compile(r'Jordan Isvy', re.I),
        re.compile(r'Carlos Eduardo Garcia Martinez', re.I),
        
        # --- Disclaimers ---
        re.compile(r'Disclosures appendix', re.I),
        re.compile(r'Copyright \d{4}', re.I),
        re.compile(r'intended for institutional investors', re.I),
        re.compile(r'If you are in scope for MiFID II', re.I),
        re.compile(r'For analyst certifications and important disclosures', re.I),
        re.compile(r'All rights reserved', re.I),
        re.compile(r'https://research.sc.com/Portal/Public/TermsConditions', re.I),
        re.compile(r'MiFID II research and inducement rules apply', re.I),
        re.compile(r'SCB accepts no liability', re.I),

        # --- Navigation / Buttons ---
        re.compile(r'^\s*Read Now\s*$', re.I),
        re.compile(r'^\s*Continue Reading\s*$', re.I),
        re.compile(r'^\s*Explore Hub\s*$', re.I),
        re.compile(r'^\s*Read the Report\s*$', re.I),
        re.compile(r'Click the button above to listen', re.I),
        re.compile(r"Can't open the report\?", re.I),
        re.compile(r'Scan code to view the report', re.I)
    ]
    
    logo_alt_patterns = [
        re.compile(r'logo', re.I),
        re.compile(r'podcast', re.I),
        re.compile(r'Listen', re.I),
        re.compile(r'barclays', re.I),
        re.compile(r'standard charte', re.I),
        re.compile(r'aranca', re.I)
    ]

    # --- Removal Pass 1: Text-based junk ---
    all_text_nodes = soup(string=True)
    nodes_to_decompose = set()

    for text_node in all_text_nodes:
        node_text = str(text_node).strip()
        if not node_text:
            continue
            
        for pattern in removal_patterns:
            if pattern.search(node_text):
                parent_to_remove = None
                current = text_node.find_parent()
                
                while current and current.name != 'body':
                    # Find the most likely "block" to remove
                    if current.name in ['tr', 'p', 'div']:
                        parent_to_remove = current
                        break
                    # If it's in a table cell, go up to the row
                    if current.name == 'td':
                         row = current.find_parent('tr')
                         if row:
                             parent_to_remove = row
                             break
                    current = current.find_parent()
                
                if parent_to_remove:
                    nodes_to_decompose.add(parent_to_remove)
                break # Move to the next text node

    # Decompose in a separate loop
    for node in nodes_to_decompose:
        if node.find_parent(): 
            node.decompose()

    # --- Removal Pass 2: Image-based junk (logos, spacers) ---
    for img_tag in soup.find_all('img'):
        alt_text = img_tag.get('alt', '')
        src_text = img_tag.get('src', '')
        is_junk = False
        
        for pattern in logo_alt_patterns:
            if pattern.search(alt_text) or pattern.search(src_text):
                is_junk = True
                break
        
        if is_junk:
            parent_to_remove = img_tag.find_parent('tr') or img_tag.find_parent('p') or img_tag
            if parent_to_remove and parent_to_remove.find_parent():
                 parent_to_remove.decompose()
            elif img_tag.find_parent():
                 img_tag.decompose()
            continue # This was junk, skip to next image
            
        # --- Process REMAINING images (Charts, etc.) ---
        src = img_tag.get('src')
        if not src:
            img_tag.decompose()
            continue
        
        # Reset styles to prevent Word from breaking layout
        for attr in ['style', 'border', 'align', 'class']:
            if img_tag.has_attr(attr):
                del img_tag[attr]
        # Set a style that Pandoc handles well
        img_tag['style'] = 'max-width: 100%; height: auto; display: block;'
        
        image_data = None
        
        # Handle Base64
        if src.startswith('data:image'):
            try:
                header, encoded = src.split(',', 1)
                image_data = base64.b64decode(encoded)
            except Exception as e:
                st.warning(f"Could not decode a Base64 image: {e}")
                img_tag.decompose()
                continue
        # Handle CID
        elif src.startswith('cid:') and msg_file_path:
            try:
                with Message(msg_file_path) as msg:
                    cid = src[4:]
                    cid_attachments = {att.cid: att for att in msg.attachments if getattr(att, 'cid', None)}
                    if cid in cid_attachments:
                        image_data = cid_attachments[cid].data
                    else:
                        img_tag.decompose()
                        continue
            except Exception as e:
                st.warning(f"Could not process CID image '{src}': {e}")
                img_tag.decompose()
                continue
        
        if image_data:
            try:
                img_type = Image.open(io.BytesIO(image_data)).format.lower() or 'png'
                img_filename = f"{uuid.uuid4()}.{img_type}"
                temp_img_path = os.path.join(temp_dir, img_filename)
                
                with open(temp_img_path, "wb") as f:
                    f.write(image_data)
                    
                img_tag['src'] = temp_img_path
            except Exception as img_e:
                st.warning(f"Failed to process and save image: {img_e}")
                img_tag.decompose()
        elif not src.startswith('data:image'):
             img_tag.decompose() # Remove broken CIDs, etc.

    # --- Fix table layouts ---
    for table in soup.find_all('table'):
        # Remove attributes that cause layout issues in Word
        for attr in ['width', 'height', 'style', 'border', 'align', 'cellpadding', 'cellspacing', 'class']:
            if table.has_attr(attr):
                del table[attr]
        # Apply a more flexible style that Pandoc understands
        table['style'] = 'border-collapse: collapse; max-width: 100%;'
        
        # Apply style to all cells
        for cell in table.find_all(['td', 'th']):
            for attr in ['style', 'border', 'align', 'class']:
                 if cell.has_attr(attr):
                    del cell[attr]
            # Add simple border and vertical alignment
            cell['style'] = 'border: 1px solid #ddd; padding: 4px; vertical-align: top;'

    # --- THIS IS THE CRITICAL LAYOUT FIX ---
    # Find the main content body.
    main_content_node = soup.find('div', class_='WordSection1') or soup.body
    
    if main_content_node:
        # Return *only the inner HTML* of the body.
        # This prevents nested <body> tags, which causes the "single column table" bug.
        return main_content_node.decode_contents()
    else:
        # Fallback for HTML fragments that don't have a <body>.
        return str(soup)


# --- MAIN UI ---
def main():
    st.markdown(f'<div class="header-container"><div class="app-title">Email Intelligence Extractor</div></div>', unsafe_allow_html=True)
    
    nav_tab1, nav_tab2, nav_tab3 = st.tabs([
        "📥 Scan & Process Emails", 
        "🔍 Query Database", 
        "📚 Manage Master Lists"
    ])

    with nav_tab1:
        st.header("Scan & Process Emails")
        scan_tab1, scan_tab2 = st.tabs(["Scan Outlook Inbox", "Upload .msg Files"])
        with scan_tab1:
            with st.container(border=True):
                st.subheader("Outlook Email Scan")
                target_email = st.text_input("Enter Mailbox Email Address:", placeholder="e.g., finance.reports@yourcompany.com", key="outlook_email_input")
                target_domain = st.text_input("Filter by Sender Domain (optional):", placeholder="e.g., jpmorgan.com", key="outlook_domain_input")
                if st.button("Scan for New Emails", type="primary", use_container_width=True, key="scan_outlook_button"):
                    if not target_email: st.warning("Please enter a mailbox email address.")
                    else:
                        with st.spinner("Authenticating and fetching emails..."):
                            token = get_graph_api_token()
                            if token:
                                emails = scan_outlook_emails(target_email, token, target_domain)
                                if emails: 
                                    # ⭐️ CHANGE: Pass no theme for Outlook scans
                                    process_emails(emails, 'outlook', email_theme=None) 
                                elif emails is not None: st.success("✅ No new unread emails found.")
        with scan_tab2:
            with st.container(border=True):
                st.subheader("Upload .msg Files")
                uploaded_files = st.file_uploader("Select .msg files to process", type=["msg"], accept_multiple_files=True, key="msg_uploader")

                # ⭐️ --- NEW CODE: Category Selection --- ⭐️
                email_categories = [
                    'Select a category...', 
                    'GCC', 
                    'Global Macro & Thematics', 
                    'Everything AI', 
                    'US Rates & FI Technicals', 
                    'EM'
                ]
                selected_email_theme = st.selectbox(
                    "Select Email Category (Optional):",
                    options=email_categories,
                    key="email_theme_select"
                )
                # ⭐️ --- END NEW CODE --- ⭐️

                if st.button("Process Uploaded Emails", type="primary", use_container_width=True, key="process_upload_button"):
                    if uploaded_files: 
                        # ⭐️ CHANGE: Pass the selected theme to the processor
                        process_emails(uploaded_files, 'upload', email_theme=selected_email_theme)
                    else: st.warning("Please upload at least one .msg file.")

    with nav_tab2:
        st.header("Query Extracted Data")

        try:
            all_data_df = query_db("SELECT * FROM email_data ORDER BY \"processedat\" DESC")
            all_data_df.columns = [x.lower() for x in all_data_df.columns]
            
            if 'processedat' in all_data_df.columns:
                all_data_df['processedat'] = pd.to_datetime(all_data_df['processedat'], errors='coerce').dt.tz_localize(None)

        except Exception as e:
            st.error(f"Could not connect to the database: {e}")
            all_data_df = pd.DataFrame()

        if all_data_df.empty:
            st.warning("The extracted data table is empty. Please process some emails first.")
        else:
            filtered_df = all_data_df.copy()

            # --- FILTER UI ---
            with st.container(border=True):
                st.subheader("📊 Filter Your Data")
                col1, col2, col3 = st.columns(3)
                with col1:
                    def get_options(column_name):
                        if column_name in all_data_df.columns:
                            return sorted([x for x in all_data_df[column_name].unique() if pd.notna(x)])
                        return []
                    selected_countries = st.multiselect("Country:", get_options('country'))
                    selected_brokers = st.multiselect("Broker Name:", get_options('brokername'))
                    selected_sectors = st.multiselect("Sector:", get_options('sector'))
                with col2:
                    selected_companies = st.multiselect("Company Name:", get_options('company'))
                    selected_content_types = st.multiselect("Email Content Type:", get_options('contenttype'))
                    
                    # ⭐️ CHANGE: Add filter for the new EmailTheme column
                    selected_email_themes = st.multiselect("Email Theme:", get_options('emailtheme'))

                with col3:
                    if 'fiscalyear' in all_data_df.columns:
                        fiscal_years = sorted([int(x) for x in all_data_df['fiscalyear'].dropna().unique()], reverse=True)
                        selected_fiscal_years = st.multiselect("Fiscal Year:", options=fiscal_years)
                    if 'fiscalquarter' in all_data_df.columns:
                        fiscal_quarters = sorted([int(x) for x in all_data_df['fiscalquarter'].dropna().unique()])
                        selected_fiscal_quarters = st.multiselect("Fiscal Quarter:", options=fiscal_quarters)
                search_query = st.text_input("Open Search (Subject or Content):", placeholder="e.g., 'acquisition' or 'earnings call'")

            # --- APPLY FILTERS ---
            if selected_countries:
                filtered_df = filtered_df[filtered_df['country'].isin(selected_countries)]
            if selected_brokers:
                filtered_df = filtered_df[filtered_df['brokername'].isin(selected_brokers)]
            if selected_sectors:
                filtered_df = filtered_df[filtered_df['sector'].isin(selected_sectors)]
            if selected_companies:
                filtered_df = filtered_df[filtered_df['company'].isin(selected_companies)]
            if selected_content_types:
                filtered_df = filtered_df[filtered_df['contenttype'].isin(selected_content_types)]
            
            # ⭐️ CHANGE: Apply EmailTheme filter
            if selected_email_themes:
                filtered_df = filtered_df[filtered_df['emailtheme'].isin(selected_email_themes)]

            if 'selected_fiscal_years' in locals() and selected_fiscal_years:
                filtered_df = filtered_df[filtered_df['fiscalyear'].isin(selected_fiscal_years)]
            if 'selected_fiscal_quarters' in locals() and selected_fiscal_quarters:
                filtered_df = filtered_df[filtered_df['fiscalquarter'].isin(selected_fiscal_quarters)]
            if search_query:
                filtered_df = filtered_df[
                    filtered_df['emailsubject'].str.contains(search_query, case=False, na=False) |
                    filtered_df['emailcontent'].str.contains(search_query, case=False, na=False)
                ]
            
            st.info(f"Displaying **{len(filtered_df)}** of **{len(all_data_df)}** total entries.")

            if not filtered_df.empty:
                temp_db_name = "financial_emails_export_filtered.db"
                if os.path.exists(temp_db_name):
                    os.remove(temp_db_name)
                conn = sqlite3.connect(temp_db_name)
                df_to_export = filtered_df.copy()
                if 'processedat' in df_to_export.columns and pd.api.types.is_datetime64_any_dtype(df_to_export['processedat']):
                    if getattr(df_to_export['processedat'].dt, 'tz', None) is not None:
                         df_to_export['processedat'] = df_to_export['processedat'].dt.tz_localize(None)
                df_to_export.to_sql('email_data', conn, if_exists='fail', index=False)
                conn.close()
                with open(temp_db_name, "rb") as fp:
                    db_bytes = fp.read()
                os.remove(temp_db_name)

                st.download_button(
                    label="Download Filtered Results (SQLite DB)",
                    data=db_bytes,
                    file_name="financial_emails_filtered.db",
                    mime="application/octet-stream"
                )
            
                # ⭐️ --- NEW WORD DOCUMENT GENERATION LOGIC --- ⭐️
                st.write("---") 
                st.subheader(f"Download Filtered Email Content ({len(filtered_df)} emails)")
                
                # ⭐️ ADDED NEW UI OPTION
                download_format = st.radio(
                    "Select download content:",
                    ["Text, Charts, and Tables", "Text Only"],
                    key="download_format_radio"
                )

                if st.button(f"Generate Word Document", key="generate_word_btn"):
                    with st.spinner("Generating Word document... This may take a moment."):
                        
                        df_for_export = filtered_df.copy()
                        if 'processedat' in df_for_export.columns:
                            df_for_export.sort_values(by='processedat', ascending=False, inplace=True)

                        try:
                            with tempfile.TemporaryDirectory() as temp_dir:
                                
                                # ⭐️ --- OPTION 1: TEXT, CHARTS, AND TABLES (PANDOC METHOD) --- ⭐️
                                if download_format == "Text, Charts, and Tables":
                                    all_html_parts = [
                                        '<!DOCTYPE html><html><head><meta charset="UTF-8">',
                                        '<style>',
                                        '@page { margin: 0.5in; size: landscape; }',
                                        'body { margin: 0.5in; font-family: "Calibri", sans-serif; font-size: 10pt; }',
                                        'img { max-width: 10in !important; height: auto !important; display: block !important; margin: 0 auto !important; page-break-inside: avoid; }',
                                        'table { width: 100% !important; max-width: 100% !important; border-collapse: collapse; }',
                                        '.email-header { background: #e8f0fe; padding: 12px; margin-bottom: 12px; border-left: 4px solid #1e70bf; }',
                                        '</style>',
                                        '</head><body>',
                                        f"<h1>Filtered Email Intelligence Report</h1>",
                                        f"<p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>",
                                        f"<p>Total Emails: {len(df_for_export)}</p>"
                                    ]

                                    for i, (index, row) in enumerate(df_for_export.iterrows()):
                                        if i > 0:
                                            all_html_parts.append('<div style="page-break-before: always;"></div>')

                                        # --- Add the standard header ---
                                        all_html_parts.append('<div class="email-header">')
                                        all_html_parts.append(f"<h2>Company: {row.get('company', 'N/A')} ({row.get('ticker', 'N/A')})</h2>")
                                        all_html_parts.append(f"<h3>Subject: {row.get('emailsubject', 'No Subject')}</h3>")
                                        date_str = row.get('processedat').strftime('%Y-%m-%d %H:%M') if pd.notna(row.get('processedat')) else 'N/A'
                                        theme_str = f" | <b>Theme:</b> {row.get('emailtheme', 'N/A')}" if pd.notna(row.get('emailtheme')) else ""
                                        all_html_parts.append(f"<p><b>Date:</b> {date_str} | <b>Broker:</b> {row.get('brokername', 'N/A')} | <b>Content Type:</b> {row.get('contenttype', 'N/A')}{theme_str}</p>")
                                        all_html_parts.append('</div>')
                                        
                                        original_html = row.get('emailcontent', '<p>No content available.</p>')
                                        
                                        # --- Get .msg file path for image extraction ---
                                        blob_name = row.get('blob_name')
                                        tmp_msg_path = None
                                        if blob_name and pd.notna(blob_name):
                                            try:
                                                blob_client = blob_service_client.get_blob_client(AZURE_CONTAINER_NAME, blob_name)
                                                msg_bytes = blob_client.download_blob().readall()
                                                fd, tmp_msg_path = tempfile.mkstemp(suffix=".msg", dir=temp_dir)
                                                with os.fdopen(fd, 'wb') as tmp_file:
                                                    tmp_file.write(msg_bytes)
                                            except Exception as e:
                                                st.warning(f"Could not retrieve .msg file {blob_name}: {e}")
                                        
                                        # ⭐️ USE THE NEW AGGRESSIVE CLEANING FUNCTION
                                        cleaned_html = clean_html_for_export(original_html, temp_dir, tmp_msg_path)
                                        all_html_parts.append(cleaned_html)

                                    all_html_parts.append('</body></html>')
                                    full_html_content = "".join(all_html_parts)

                                    output_filename = os.path.join(temp_dir, "output.docx")
                                    
                                    pypandoc.convert_text(
                                        full_html_content,
                                        'docx',
                                        format='html',
                                        outputfile=output_filename,
                                        extra_args=[
                                            f'--resource-path={temp_dir}',
                                            '--wrap=none',
                                            '--standalone',
                                            '-V', 'geometry:landscape',
                                            '-V', 'geometry:margin=0.5in' # Using 0.5in margin
                                        ]
                                    )
                                    
                                    with open(output_filename, "rb") as f:
                                        output_docx_bytes = f.read()

                                # ⭐️ --- OPTION 2: TEXT ONLY (PYTHON-DOCX METHOD) --- ⭐️
                                else: # download_format == "Text Only"
                                    document = docx.Document()
                                    doc_io = io.BytesIO()

                                    for i, (index, row) in enumerate(df_for_export.iterrows()):
                                        if i > 0:
                                            document.add_page_break()
                                        
                                        # --- Add the standard header ---
                                        document.add_heading(f"Company: {row.get('company', 'N/A')} ({row.get('ticker', 'N/A')})", level=2)
                                        document.add_heading(f"Subject: {row.get('emailsubject', 'No Subject')}", level=3)
                                        date_str = row.get('processedat').strftime('%Y-%m-%d %H:%M') if pd.notna(row.get('processedat')) else 'N/A'
                                        theme_str = f" | Theme: {row.get('emailtheme', 'N/A')}" if pd.notna(row.get('emailtheme')) else ""
                                        p = document.add_paragraph()
                                        p.add_run(f"Date: {date_str} | Broker: {row.get('brokername', 'N/A')}{theme_str}").bold = True
                                        
                                        # --- Clean and add text content ---
                                        original_html = row.get('emailcontent', '')
                                        
                                        # Use the *same* aggressive cleaning function
                                        # We don't need temp_dir or msg_path since we're stripping images anyway
                                        cleaned_html = clean_html_for_export(original_html, temp_dir, None) 
                                        
                                        soup = BeautifulSoup(cleaned_html, 'html.parser')
                                        body_text = soup.get_text(separator='\n', strip=True)
                                        
                                        # Add the cleaned text
                                        document.add_paragraph(body_text)

                                    document.save(doc_io)
                                    output_docx_bytes = doc_io.getvalue()

                            # --- Set session state for download ---
                            st.session_state['word_data'] = output_docx_bytes
                            st.session_state['word_filename'] = f"email_content_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            st.success("✅ Word document generated successfully!")

                        except Exception as e:
                            st.error(f"Failed to generate Word document. Error: {e}")
                            # For debugging, you can write the problematic HTML to a file
                            if 'full_html_content' in locals():
                                with open(os.path.join(temp_dir, "debug.html"), "w", encoding="utf-8") as f:
                                    f.write(full_html_content)
                                st.info("A 'debug.html' file has been saved in the temp directory for inspection.")

                        st.rerun()

                if st.session_state.get('word_data'):
                    st.download_button(
                        label=f"✅ Click to Download: {st.session_state['word_filename']}",
                        data=st.session_state['word_data'],
                        file_name=st.session_state['word_filename'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                        on_click=lambda: st.session_state.pop('word_data', None)
                    )

            # --- DISPLAY DATAFRAME with Download Link ---
            st.write("---")
            if 'blob_name' in filtered_df.columns:
                display_df = filtered_df.copy()
                display_df['download_link'] = display_df['blob_name'].apply(
                    lambda name: generate_sas_url(AZURE_CONTAINER_NAME, name)
                )
                st.dataframe(
                    display_df,
                    column_config={
                        "download_link": st.column_config.LinkColumn(
                            "Original Email",
                            help="Click to download the original .msg file (link expires in 1 hour)",
                            display_text="⬇️ Download"
                        ),
                        "emailcontent": None,
                    },
                    use_container_width=True
                )
            else:
                st.dataframe(filtered_df, use_container_width=True)

    with nav_tab3:
        st.header("Manage Master Lists")
        st.info(f"This data is read from '{MASTER_DB_NAME}'. To update it, please use your external import script.")
        st.subheader("Master Company List")
        st.dataframe(get_master_company_data(), use_container_width=True)
        st.subheader("Master Broker List")
        broker_names = get_master_broker_names()
        if broker_names:
            st.dataframe(pd.DataFrame(broker_names, columns=["Broker Name"]), use_container_width=True)
        else:
            st.warning(f"The 'master_brokers' table was not found in '{MASTER_DB_NAME}'.")

if __name__ == "__main__":
    main()